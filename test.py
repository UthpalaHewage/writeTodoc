from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading("Adverbs", 0)

document.add_heading('Sub Heading 01', level=1)


p = document.add_paragraph('content for sub heading 01 with styles.')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

# document.add_page_break()
document.add_picture('test.jpg', width=Inches(5),height=Inches(3))

document.add_heading('Sub Heading 02', level=1)

p = document.add_paragraph('content for sub heading 02 with styles.')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True







document.save('demo.docx')
